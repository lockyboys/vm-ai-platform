"""SPS PDF/image upload demonstration web application."""
from __future__ import annotations

import os
import secrets
import tempfile
from functools import wraps
from pathlib import Path
from typing import Callable, ParamSpec, TypeVar

from docx import Document
from flask import Flask, abort, flash, redirect, render_template_string, request, send_file, session, url_for
from werkzeug.utils import secure_filename

from common.database import CommonDatabase
from engine.processor.work.file_work_service import DEFAULT_WORK_OUTPUT_ROOT, FileWorkService

P = ParamSpec("P")
R = TypeVar("R")
OUTPUT_ROOT = Path(DEFAULT_WORK_OUTPUT_ROOT).resolve()

LOGIN_TEMPLATE = """<!doctype html><html lang="ko"><meta charset="utf-8">
<title>SPS Document Intelligence</title><style>
body{margin:0;font-family:"Malgun Gothic",sans-serif;background:#f4f7fb;color:#182033}main{width:min(440px,calc(100% - 40px));margin:12vh auto;background:#fff;padding:42px;border-radius:16px;box-shadow:0 18px 48px #23396c1f}h1{margin:0 0 8px;font-size:26px}.sub{color:#63708a;line-height:1.6}label{display:block;font-weight:700;margin-top:20px}input{box-sizing:border-box;width:100%;margin-top:8px;padding:12px;border:1px solid #ccd5e5;border-radius:8px;font-size:16px}button{width:100%;margin-top:28px;padding:13px;border:0;border-radius:8px;background:#1d4ed8;color:#fff;font-weight:700;font-size:16px}.notice{margin:18px 0;padding:12px;border-radius:8px;background:#fff4e5;color:#8a4b00}</style>
<main><h1>SPS Document Intelligence</h1><p class="sub">로그인 후 PDF 또는 이미지 파일을 올리면 SPS가 내용을 읽고 Repository에 기록한 뒤 공식 리포트를 생성합니다.</p>{% with messages=get_flashed_messages() %}{% if messages %}<p class="notice">{{messages[-1]}}</p>{% endif %}{% endwith %}<form method="post"><input type="hidden" name="csrf_token" value="{{csrf_token}}"><label>사용자 ID<input name="username" required></label><label>비밀번호<input type="password" name="password" required></label><button>로그인</button></form></main></html>"""

EDITOR_TEMPLATE = """<!doctype html><html lang="ko"><meta charset="utf-8">
<title>SPS 파일 처리</title><style>
body{margin:0;font-family:"Malgun Gothic",sans-serif;background:#f4f7fb;color:#182033}header{padding:18px 5%;background:#102a63;color:#fff;display:flex;justify-content:space-between}main{max-width:840px;margin:34px auto;padding:0 24px}.card{background:#fff;padding:30px;border-radius:16px;box-shadow:0 14px 38px #23396c14}h1{margin-top:0}.guide{color:#63708a;line-height:1.8}.flow{background:#eef4ff;padding:16px;border-radius:10px;font-weight:700}.flash{padding:12px;background:#fff4e5;color:#8a4b00;border-radius:8px}input{box-sizing:border-box;width:100%;margin-top:10px;padding:12px;border:1px solid #ccd5e5;border-radius:8px;font:inherit}button{margin-top:26px;padding:13px 24px;background:#1d4ed8;color:#fff;border:0;border-radius:8px;font-weight:700;font-size:16px}a{color:#fff}</style>
<header><strong>SPS Document Intelligence · File Work</strong><span>{{username}} · <a href="{{url_for('logout')}}">로그아웃</a></span></header>
<main><div class="card"><h1>PDF·이미지 파일 처리</h1><p class="flow">로그인 → 업로드 → PDF 텍스트 추출/OCR → Repository DB 저장 → DOCX·실행 리포트 다운로드</p><p class="guide">지원 형식: PDF, JPG, JPEG, PNG. 업로드한 원본·추출 텍스트·DOCX·Markdown 리포트는 하나의 Work Session에 연결됩니다.</p>{% with messages=get_flashed_messages() %}{% if messages %}<p class="flash">{{messages[-1]}}</p>{% endif %}{% endwith %}<form method="post" action="{{url_for('process_file')}}" enctype="multipart/form-data"><input type="hidden" name="csrf_token" value="{{csrf_token}}"><label><b>원본 파일 선택</b><input type="file" name="source_file" accept=".pdf,.jpg,.jpeg,.png" required></label><button type="submit">파일 읽기·DB 저장·리포트 생성</button></form></div></main></html>"""

RESULT_TEMPLATE = """<!doctype html><html lang="ko"><meta charset="utf-8">
<title>SPS 파일 처리 결과</title><style>
body{margin:0;font-family:"Malgun Gothic",sans-serif;background:#f4f7fb;color:#182033}header{padding:18px 5%;background:#102a63;color:#fff}main{max-width:820px;margin:34px auto;padding:0 24px}.card{background:#fff;padding:28px;border-radius:16px;box-shadow:0 14px 38px #23396c14}.ok{color:#087443;font-weight:700}.buttons a,.buttons button{display:inline-block;margin:10px 10px 0 0;padding:12px 18px;border:0;border-radius:8px;background:#1d4ed8;color:#fff;text-decoration:none;font:700 16px "Malgun Gothic",sans-serif;cursor:pointer}.buttons .secondary{background:#556274}li{margin:10px 0}code{word-break:break-all}@media print{body{background:#fff}.buttons{display:none}header{padding:0 0 16px;background:#fff;color:#182033}.card{padding:0;box-shadow:none}main{margin:0;max-width:none}}</style>
<header><strong>SPS Document Intelligence · Result</strong></header><main><div class="card"><h1>파일 처리 및 Repository 저장 완료</h1><p class="ok">원본 업로드 → 텍스트 추출/OCR → sp_work_session → sp_work_item → sp_work_asset → DOCX·리포트</p><h2>공식 저장 식별자</h2><ul><li>Work Session ID: <code>{{work_session_id}}</code></li><li>Work Item ID: <code>{{work_item_id}}</code></li><li>Source Object ID: <code>{{source_object_id}}</code></li><li>추출 텍스트 길이: {{extracted_text_length}}</li></ul><h2>저장 검증</h2><ul><li>sp_work_session 1건 저장</li><li>sp_work_item 1건 저장</li><li>sp_work_asset 4건 저장: 원본·추출 텍스트·DOCX·Markdown</li></ul><div class="buttons">{% if demo_report_available %}<a href="{{url_for('preview_demo_report')}}">{{demo_report_title}} 출력</a>{% endif %}<a href="{{url_for('preview_docx', work_session_id=work_session_id)}}">이번 DOCX 미리보기·출력</a><a href="{{url_for('download', work_session_id=work_session_id, artifact='docx')}}">DOCX 다운로드</a><a href="{{url_for('download', work_session_id=work_session_id, artifact='report')}}">실행 리포트 다운로드</a><button type="button" onclick="window.print()">결과 화면 출력</button><a class="secondary" href="{{url_for('editor')}}">다른 파일 처리</a></div></div></main></html>"""

PRINT_TEMPLATE = """<!doctype html><html lang="ko"><meta charset="utf-8"><title>{{report_title}}</title><style>body{margin:0;font-family:"Malgun Gothic",sans-serif;background:#f4f7fb;color:#182033}.toolbar{padding:16px 5%;background:#102a63;color:#fff;display:flex;gap:10px;align-items:center}.toolbar button,.toolbar a{padding:10px 14px;border:0;border-radius:8px;background:#fff;color:#102a63;text-decoration:none;font:700 15px "Malgun Gothic",sans-serif;cursor:pointer}.document{max-width:820px;margin:34px auto;background:#fff;padding:56px 64px;box-shadow:0 14px 38px #23396c14}.document h1{font-size:28px}.document h2{margin-top:28px;font-size:21px}.document p{line-height:1.8;white-space:pre-wrap}.document table{border-collapse:collapse;width:100%;margin:18px 0}.document th,.document td{border:1px solid #b9c3d3;padding:8px;text-align:left;vertical-align:top}@media print{body{background:#fff}.toolbar{display:none}.document{margin:0;max-width:none;padding:0;box-shadow:none}}</style><div class="toolbar"><strong>{{report_title}}</strong><button type="button" onclick="window.print()">이 리포트 출력</button>{% if download_url %}<a href="{{download_url}}">DOCX 다운로드</a>{% endif %}<a href="{{url_for('editor')}}">파일 처리로 돌아가기</a></div><article class="document">{% for block in blocks %}{% if block.kind == 'title' %}<h1>{{block.text}}</h1>{% elif block.kind == 'heading' %}<h2>{{block.text}}</h2>{% elif block.kind == 'table' %}<table>{% for row in block.rows %}<tr>{% for cell in row %}<td>{{cell}}</td>{% endfor %}</tr>{% endfor %}</table>{% else %}<p>{{block.text}}</p>{% endif %}{% endfor %}</article></html>"""


def _required_setting(name: str) -> str:
    value = os.getenv(name, "").strip()
    if not value:
        raise RuntimeError(f"{name} environment variable is required.")
    return value


def _csrf_token() -> str:
    token = session.get("csrf_token")
    if token is None:
        token = secrets.token_urlsafe(32)
        session["csrf_token"] = token
    return token


def _validate_csrf() -> None:
    if not secrets.compare_digest(request.form.get("csrf_token", ""), session.get("csrf_token", "")):
        abort(400, "Invalid CSRF token.")


def _read_docx_blocks(docx_path: Path) -> list[dict[str, object]]:
    report_document = Document(docx_path)
    blocks: list[dict[str, object]] = []
    for paragraph in report_document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        kind = "title" if style_name == "Title" else "heading" if style_name.startswith("Heading") else "paragraph"
        blocks.append({"kind": kind, "text": text})
    for table in report_document.tables:
        blocks.append({"kind": "table", "rows": [[cell.text for cell in row.cells] for row in table.rows]})
    return blocks


def _configured_demo_report() -> tuple[str, Path] | None:
    report_title = os.getenv("SPS_DEMO_REPORT_TITLE", "").strip()
    work_asset_id = os.getenv("SPS_DEMO_REPORT_WORK_ASSET_ID", "").strip()
    if not report_title or not work_asset_id:
        return None
    database = CommonDatabase(database_role="STORY")
    try:
        row = database.fetch_one(
            """
            SELECT asset_name, asset_path
            FROM sp_work_asset
            WHERE work_asset_id = %s
              AND asset_type_code = 'DOCX_REPORT'
              AND asset_status_code = 'STORED'
              AND deleted_dt IS NULL
            """,
            (work_asset_id,),
        )
    finally:
        database.close()
    if not row:
        return None
    report_path = Path(str(row["asset_path"])).resolve()
    if OUTPUT_ROOT not in report_path.parents or not report_path.is_file():
        return None
    return report_title, report_path


def _require_login(view: Callable[P, R]) -> Callable[P, R]:
    @wraps(view)
    def wrapped(*args: P.args, **kwargs: P.kwargs) -> R:
        if not session.get("document_user"):
            return redirect(url_for("login"))
        return view(*args, **kwargs)
    return wrapped


def create_app() -> Flask:
    app = Flask(__name__)
    app.config["SECRET_KEY"] = _required_setting("SPS_DOCUMENT_DEMO_SECRET_KEY")
    app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024
    service = FileWorkService(output_root=OUTPUT_ROOT)

    @app.get("/")
    def index():
        return redirect(url_for("editor"))

    @app.route("/login", methods=["GET", "POST"])
    def login():
        if request.method == "POST":
            _validate_csrf()
            if (
                secrets.compare_digest(request.form.get("username", "").strip(), _required_setting("SPS_DOCUMENT_DEMO_USERNAME"))
                and secrets.compare_digest(request.form.get("password", ""), _required_setting("SPS_DOCUMENT_DEMO_PASSWORD"))
            ):
                session.clear()
                session["document_user"] = request.form["username"].strip()
                _csrf_token()
                return redirect(url_for("editor"))
            flash("로그인 정보를 확인해 주세요.")
        return render_template_string(LOGIN_TEMPLATE, csrf_token=_csrf_token())

    @app.get("/logout")
    def logout():
        session.clear()
        return redirect(url_for("login"))

    @app.get("/editor")
    @_require_login
    def editor():
        return render_template_string(EDITOR_TEMPLATE, username=session["document_user"], csrf_token=_csrf_token())

    @app.post("/process-file")
    @_require_login
    def process_file():
        _validate_csrf()
        uploaded = request.files.get("source_file")
        filename = secure_filename(uploaded.filename or "") if uploaded else ""
        if not uploaded or not filename:
            flash("업로드할 파일을 선택해 주세요.")
            return redirect(url_for("editor"))
        suffix = Path(filename).suffix.lower()
        if suffix not in {".pdf", ".jpg", ".jpeg", ".png"}:
            flash("PDF, JPG, JPEG, PNG 파일만 업로드할 수 있습니다.")
            return redirect(url_for("editor"))
        with tempfile.TemporaryDirectory(prefix="sps_file_work_") as directory:
            source_path = Path(directory) / filename
            uploaded.save(source_path)
            try:
                result = service.process(
                    upload_path=source_path,
                    requested_by=session["document_user"],
                    client_ip=request.remote_addr or "",
                )
            except Exception as error:
                app.logger.exception("File work processing failed")
                flash(f"파일 처리에 실패했습니다: {error}")
                return redirect(url_for("editor"))
        artifacts = session.get("work_artifacts", {})
        artifacts[result.work_session_id] = {"docx": str(result.docx_path), "report": str(result.report_path)}
        session["work_artifacts"] = artifacts
        return render_template_string(
            RESULT_TEMPLATE,
            work_session_id=result.work_session_id,
            work_item_id=result.work_item_id,
            source_object_id=result.source_object_id,
            extracted_text_length=result.extracted_text_length,
            demo_report_available=False,
            demo_report_title=os.getenv("SPS_DEMO_REPORT_TITLE", "").strip(),
        )


    @app.get("/preview/<work_session_id>/docx")
    @_require_login
    def preview_docx(work_session_id: str):
        artifact_path = session.get("work_artifacts", {}).get(work_session_id, {}).get("docx")
        if not artifact_path:
            abort(404)
        resolved_path = Path(artifact_path).resolve()
        if OUTPUT_ROOT not in resolved_path.parents or not resolved_path.is_file():
            abort(404)
        return render_template_string(
            PRINT_TEMPLATE,
            report_title="이번 파일 처리 DOCX 리포트",
            download_url=url_for("download", work_session_id=work_session_id, artifact="docx"),
            blocks=_read_docx_blocks(resolved_path),
        )

    @app.get("/report/demo")
    @_require_login
    def preview_demo_report():
        configured_report = _configured_demo_report()
        if configured_report is None:
            abort(404)
        report_title, report_path = configured_report
        blocks = _read_docx_blocks(report_path)
        for block in blocks:
            if block["kind"] == "title":
                block["text"] = report_title
                break
        else:
            blocks.insert(0, {"kind": "title", "text": report_title})
        return render_template_string(
            PRINT_TEMPLATE,
            report_title=report_title,
            download_url=None,
            blocks=blocks,
        )

    @app.get("/download/<work_session_id>/<artifact>")
    @_require_login
    def download(work_session_id: str, artifact: str):
        if artifact not in {"docx", "report"}:
            abort(404)
        artifact_path = session.get("work_artifacts", {}).get(work_session_id, {}).get(artifact)
        if not artifact_path:
            abort(404)
        resolved_path = Path(artifact_path).resolve()
        if OUTPUT_ROOT not in resolved_path.parents or not resolved_path.is_file():
            abort(404)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if artifact == "docx" else "text/markdown; charset=utf-8"
        return send_file(resolved_path, as_attachment=True, download_name=resolved_path.name, mimetype=mimetype)

    return app


if __name__ == "__main__":
    create_app().run(host=os.getenv("SPS_DOCUMENT_DEMO_HOST", "127.0.0.1"), port=int(os.getenv("SPS_DOCUMENT_DEMO_PORT", "8010")), debug=False)
