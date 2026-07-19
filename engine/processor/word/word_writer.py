"""
SPS Word Writer

Purpose:
    DocumentModel을 Microsoft Word DOCX 파일로 출력한다.

Principles:
    - Repository First
    - Generator First
    - Metadata Driven
    - Single Source of Truth
    - No Hardcoding
"""

from __future__ import annotations

from pathlib import Path
from typing import Final

from docx import Document
from docx.document import Document as WordDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from engine.processor.word.document_model import (
    DocumentHeading,
    DocumentModel,
    DocumentParagraph,
    DocumentTable,
)


class WordWriter:
    """DocumentModel을 DOCX Artifact로 출력한다."""

    DEFAULT_FONT_NAME: Final[str] = "맑은 고딕"
    DEFAULT_FONT_SIZE: Final[int] = 10
    TITLE_FONT_SIZE: Final[int] = 20
    TABLE_FONT_SIZE: Final[int] = 8

    def __init__(
        self,
        *,
        font_name: str = DEFAULT_FONT_NAME,
        font_size: int = DEFAULT_FONT_SIZE,
        title_font_size: int = TITLE_FONT_SIZE,
        table_font_size: int = TABLE_FONT_SIZE,
    ) -> None:
        self.font_name = self._require_text(
            font_name,
            "font_name",
        )
        self.font_size = self._require_positive_integer(
            font_size,
            "font_size",
        )
        self.title_font_size = self._require_positive_integer(
            title_font_size,
            "title_font_size",
        )
        self.table_font_size = self._require_positive_integer(
            table_font_size,
            "table_font_size",
        )

    def save(
        self,
        model: DocumentModel,
        output_path: str | Path,
    ) -> Path:
        """
        DocumentModel을 DOCX 파일로 저장한다.

        Args:
            model:
                출력할 DocumentModel.
            output_path:
                저장할 DOCX 파일 경로.

        Returns:
            생성된 DOCX 파일의 절대 경로.
        """

        if not isinstance(model, DocumentModel):
            raise TypeError(
                "model must be an instance of DocumentModel."
            )

        resolved_output_path = self._resolve_output_path(
            output_path
        )

        document = Document()

        self._configure_document(document)
        self._configure_styles(document)
        self._write_document_properties(document, model)
        self._write_title(document, model)
        self._write_blocks(document, model)

        document.save(str(resolved_output_path))

        if not resolved_output_path.exists():
            raise RuntimeError(
                "Word file was not created: "
                f"{resolved_output_path}"
            )

        return resolved_output_path

    def _resolve_output_path(
        self,
        output_path: str | Path,
    ) -> Path:
        """출력 경로를 검증하고 상위 디렉터리를 생성한다."""

        if isinstance(output_path, str):
            output_path = output_path.strip()

        if not output_path:
            raise ValueError(
                "output_path must not be empty."
            )

        resolved_path = Path(output_path).expanduser()

        if resolved_path.suffix.lower() != ".docx":
            resolved_path = resolved_path.with_suffix(".docx")

        resolved_path = resolved_path.resolve()
        resolved_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        return resolved_path

    def _configure_document(
        self,
        document: WordDocument,
    ) -> None:
        """Word 문서의 기본 Page 설정을 적용한다."""

        for section in document.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

    def _configure_styles(
        self,
        document: WordDocument,
    ) -> None:
        """Word 문서의 기본 Style을 설정한다."""

        normal_style = document.styles["Normal"]
        normal_style.font.name = self.font_name
        normal_style.font.size = Pt(self.font_size)

        self._set_style_east_asia_font(
            normal_style,
            self.font_name,
        )

        for level in range(1, 10):
            style_name = f"Heading {level}"

            if style_name not in document.styles:
                continue

            heading_style = document.styles[style_name]
            heading_style.font.name = self.font_name

            self._set_style_east_asia_font(
                heading_style,
                self.font_name,
            )

    def _write_document_properties(
        self,
        document: WordDocument,
        model: DocumentModel,
    ) -> None:
        """DOCX Core Property를 설정한다."""

        properties = document.core_properties

        properties.title = model.title
        properties.subject = (
            model.metadata.description or model.title
        )
        properties.author = (
            model.metadata.processor_code
        )
        properties.keywords = ", ".join(
            [
                model.metadata.artifact_type_code,
                model.metadata.source_object_code,
                model.metadata.processor_code,
            ]
        )
        properties.comments = (
            "Generated by Story Programming Framework."
        )
        properties.created = model.metadata.generated_dt
        properties.modified = model.metadata.generated_dt

    def _write_title(
        self,
        document: WordDocument,
        model: DocumentModel,
    ) -> None:
        """문서 제목을 출력한다."""

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(model.title)
        run.bold = True
        run.font.name = self.font_name
        run.font.size = Pt(self.title_font_size)

        self._set_run_east_asia_font(
            run,
            self.font_name,
        )

        paragraph.space_after = Pt(12)

    def _write_blocks(
        self,
        document: WordDocument,
        model: DocumentModel,
    ) -> None:
        """Document Block을 순서대로 출력한다."""

        for block in model.blocks:
            if isinstance(block, DocumentHeading):
                self._write_heading(
                    document,
                    block,
                )
                continue

            if isinstance(block, DocumentParagraph):
                self._write_paragraph(
                    document,
                    block,
                )
                continue

            if isinstance(block, DocumentTable):
                self._write_table(
                    document,
                    block,
                )
                continue

            raise TypeError(
                "Unsupported document block type: "
                f"{type(block).__name__}"
            )

    def _write_heading(
        self,
        document: WordDocument,
        block: DocumentHeading,
    ) -> None:
        """Heading Block을 출력한다."""

        paragraph = document.add_heading(
            block.text,
            level=block.level,
        )

        for run in paragraph.runs:
            run.font.name = self.font_name
            self._set_run_east_asia_font(
                run,
                self.font_name,
            )

    def _write_paragraph(
        self,
        document: WordDocument,
        block: DocumentParagraph,
    ) -> None:
        """Paragraph Block을 출력한다."""

        paragraph = document.add_paragraph()
        paragraph.alignment = self._resolve_alignment(
            block.alignment
        )

        run = paragraph.add_run(block.text)
        run.bold = block.bold
        run.italic = block.italic
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)

        self._set_run_east_asia_font(
            run,
            self.font_name,
        )

    def _write_table(
        self,
        document: WordDocument,
        block: DocumentTable,
    ) -> None:
        """Table Block을 출력한다."""

        if block.title:
            title_paragraph = document.add_paragraph()

            title_run = title_paragraph.add_run(
                block.title
            )
            title_run.bold = True
            title_run.font.name = self.font_name
            title_run.font.size = Pt(
                self.font_size + 1
            )

            self._set_run_east_asia_font(
                title_run,
                self.font_name,
            )

        column_count = len(block.headers)
        row_count = len(block.rows) + 1

        table = document.add_table(
            rows=row_count,
            cols=column_count,
        )

        table.style = block.style_name
        table.autofit = block.autofit

        self._write_table_headers(
            table,
            block,
        )
        self._write_table_rows(
            table,
            block,
        )

        if block.repeat_header:
            self._set_repeat_table_header(
                table.rows[0]
            )

        document.add_paragraph()

    def _write_table_headers(
        self,
        table,
        block: DocumentTable,
    ) -> None:
        """Table Header를 출력한다."""

        header_row = table.rows[0]

        for column_index, header in enumerate(
            block.headers
        ):
            cell = header_row.cells[column_index]
            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            run = paragraph.add_run(header)
            run.bold = True
            run.font.name = self.font_name
            run.font.size = Pt(self.table_font_size)

            self._set_run_east_asia_font(
                run,
                self.font_name,
            )

            self._set_cell_background(
                cell,
                "D9EAF7",
            )

    def _write_table_rows(
        self,
        table,
        block: DocumentTable,
    ) -> None:
        """Table Data Row를 출력한다."""

        for row_index, row_values in enumerate(
            block.rows,
            start=1,
        ):
            table_row = table.rows[row_index]

            for column_index, value in enumerate(
                row_values
            ):
                cell = table_row.cells[column_index]
                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

                paragraph = cell.paragraphs[0]
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                )

                run = paragraph.add_run(value)
                run.font.name = self.font_name
                run.font.size = Pt(
                    self.table_font_size
                )

                self._set_run_east_asia_font(
                    run,
                    self.font_name,
                )

    def _resolve_alignment(
        self,
        alignment: str,
    ) -> WD_ALIGN_PARAGRAPH:
        """문자열 Alignment를 python-docx Enum으로 변환한다."""

        alignment_map = {
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
            "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        try:
            return alignment_map[alignment.upper()]
        except KeyError as exc:
            raise ValueError(
                "Unsupported paragraph alignment: "
                f"{alignment}"
            ) from exc

    @staticmethod
    def _set_repeat_table_header(row) -> None:
        """표 Header가 다음 Page에서도 반복되도록 설정한다."""

        table_row_properties = (
            row._tr.get_or_add_trPr()
        )

        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(
            qn("w:val"),
            "true",
        )

        table_row_properties.append(
            repeat_header
        )

    @staticmethod
    def _set_cell_background(
        cell,
        fill: str,
    ) -> None:
        """Table Cell 배경색을 설정한다."""

        cell_properties = (
            cell._tc.get_or_add_tcPr()
        )

        shading = cell_properties.find(
            qn("w:shd")
        )

        if shading is None:
            shading = OxmlElement("w:shd")
            cell_properties.append(shading)

        shading.set(
            qn("w:fill"),
            fill,
        )

    @staticmethod
    def _set_run_east_asia_font(
        run,
        font_name: str,
    ) -> None:
        """Run에 한글 Font를 적용한다."""

        run.font.name = font_name

        run_properties = (
            run._element.get_or_add_rPr()
        )
        run_fonts = run_properties.rFonts

        if run_fonts is None:
            run_fonts = OxmlElement("w:rFonts")
            run_properties.append(run_fonts)

        run_fonts.set(
            qn("w:ascii"),
            font_name,
        )
        run_fonts.set(
            qn("w:hAnsi"),
            font_name,
        )
        run_fonts.set(
            qn("w:eastAsia"),
            font_name,
        )

    @staticmethod
    def _set_style_east_asia_font(
        style,
        font_name: str,
    ) -> None:
        """Style에 한글 Font를 적용한다."""

        style.font.name = font_name

        style_properties = (
            style.element.get_or_add_rPr()
        )
        run_fonts = style_properties.rFonts

        if run_fonts is None:
            run_fonts = OxmlElement("w:rFonts")
            style_properties.append(run_fonts)

        run_fonts.set(
            qn("w:ascii"),
            font_name,
        )
        run_fonts.set(
            qn("w:hAnsi"),
            font_name,
        )
        run_fonts.set(
            qn("w:eastAsia"),
            font_name,
        )

    @staticmethod
    def _require_text(
        value: str,
        field_name: str,
    ) -> str:
        """필수 문자열 값을 검증한다."""

        normalized_value = str(value).strip()

        if not normalized_value:
            raise ValueError(
                f"{field_name} must not be empty."
            )

        return normalized_value

    @staticmethod
    def _require_positive_integer(
        value: int,
        field_name: str,
    ) -> int:
        """양의 정수 값을 검증한다."""

        if isinstance(value, bool):
            raise TypeError(
                f"{field_name} must be an integer."
            )

        try:
            normalized_value = int(value)
        except (TypeError, ValueError) as exc:
            raise TypeError(
                f"{field_name} must be an integer."
            ) from exc

        if normalized_value <= 0:
            raise ValueError(
                f"{field_name} must be greater than zero."
            )

        return normalized_value
