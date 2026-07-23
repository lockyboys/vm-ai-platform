"""SPS interactive document generation service.

This service turns authenticated user input into a neutral DocumentModel and
exports the model through the existing WordWriter. It owns no database state.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from re import sub
from typing import Final
from uuid import uuid4
from zoneinfo import ZoneInfo

from docx import Document

from engine.processor.word.document_model import (
    DocumentHeading,
    DocumentMetadata,
    DocumentModel,
)
from engine.processor.word.word_writer import WordWriter


SEOUL_TIME_ZONE: Final[str] = "Asia/Seoul"
DEFAULT_OUTPUT_ROOT: Final[Path] = Path("output/document_generation")


@dataclass(frozen=True, slots=True)
class DocumentGenerationRequest:
    """Validated input received from the document editor."""

    title: str
    part_title: str
    philosophy: str
    chapter_title: str
    section_title: str
    body: str

    @classmethod
    def from_mapping(cls, values: dict[str, str]) -> "DocumentGenerationRequest":
        normalized = {
            key: str(values.get(key, "")).strip()
            for key in (
                "title",
                "part_title",
                "philosophy",
                "chapter_title",
                "section_title",
                "body",
            )
        }
        missing = [key for key, value in normalized.items() if not value]
        if missing:
            raise ValueError(
                "Required document input is missing: " + ", ".join(missing)
            )
        return cls(**normalized)


@dataclass(frozen=True, slots=True)
class DocumentGenerationResult:
    """Artifacts and verification result returned to the web layer."""

    request_id: str
    document: DocumentModel
    docx_path: Path
    report_path: Path
    headings: tuple[str, ...]


class DocumentGenerationService:
    """Generate a document artifact from user input."""

    def __init__(
        self,
        *,
        output_root: str | Path = DEFAULT_OUTPUT_ROOT,
        writer: WordWriter | None = None,
    ) -> None:
        self.output_root = Path(output_root)
        self.writer = writer or WordWriter()

    def generate(
        self,
        *,
        request: DocumentGenerationRequest,
        requested_by: str,
    ) -> DocumentGenerationResult:
        """Build, export, read back, and report one document generation run."""

        request_id = uuid4().hex
        document = self._build_document(
            request=request,
            requested_by=requested_by,
            request_id=request_id,
        )
        artifact_directory = self.output_root / request_id
        artifact_directory.mkdir(parents=True, exist_ok=True)

        document_name = self._file_stem(request.title)
        docx_path = self.writer.save(
            model=document,
            output_path=artifact_directory / f"{document_name}.docx",
        )
        headings = tuple(self._read_exported_headings(docx_path))
        expected_headings = tuple(
            block.text
            for block in document.blocks
            if isinstance(block, DocumentHeading)
        )
        if headings != expected_headings:
            raise RuntimeError(
                "Exported DOCX heading verification failed. "
                f"expected={expected_headings!r}, actual={headings!r}"
            )

        report_path = self._write_report(
            request_id=request_id,
            document=document,
            requested_by=requested_by,
            docx_path=docx_path,
            headings=headings,
        )
        return DocumentGenerationResult(
            request_id=request_id,
            document=document,
            docx_path=docx_path,
            report_path=report_path,
            headings=headings,
        )

    @staticmethod
    def _build_document(
        *,
        request: DocumentGenerationRequest,
        requested_by: str,
        request_id: str,
    ) -> DocumentModel:
        metadata = DocumentMetadata(
            document_name=DocumentGenerationService._file_stem(request.title),
            source_object_code="DOCUMENT",
            description="사용자 입력으로 생성한 SPS Document Framework DOCX 산출물",
            attributes={
                "requested_by": requested_by,
                "request_id": request_id,
                "generation_channel": "DOCUMENT_WEB",
            },
        )
        document = DocumentModel(title=request.title, metadata=metadata)
        document.add_numbered_heading(
            request.part_title,
            level=1,
            heading_label="Part",
        )
        document.add_paragraph(request.philosophy, italic=True)
        document.add_numbered_heading(
            request.chapter_title,
            level=2,
            heading_label="Chapter",
        )
        document.add_numbered_heading(request.section_title, level=3)
        for paragraph in DocumentGenerationService._paragraphs(request.body):
            document.add_paragraph(paragraph)
        return document

    @staticmethod
    def _paragraphs(body: str) -> tuple[str, ...]:
        paragraphs = tuple(
            line.strip()
            for line in body.replace("\r\n", "\n").split("\n")
            if line.strip()
        )
        if not paragraphs:
            raise ValueError("body must contain at least one paragraph.")
        return paragraphs

    @staticmethod
    def _file_stem(value: str) -> str:
        normalized = sub(r"[^0-9A-Za-z가-힣._-]+", "_", value.strip())
        normalized = normalized.strip("._")
        return normalized[:80] or "sps_document"

    @staticmethod
    def _read_exported_headings(docx_path: Path) -> list[str]:
        exported = Document(docx_path)
        return [
            paragraph.text.strip()
            for paragraph in exported.paragraphs
            if paragraph.style.name.startswith("Heading")
            and paragraph.text.strip()
        ]

    @staticmethod
    def _write_report(
        *,
        request_id: str,
        document: DocumentModel,
        requested_by: str,
        docx_path: Path,
        headings: tuple[str, ...],
    ) -> Path:
        report_path = docx_path.with_suffix(".md")
        generated_dt = datetime.now(
            ZoneInfo(SEOUL_TIME_ZONE)
        ).strftime("%Y-%m-%d %H:%M:%S KST")
        report_lines = [
            "# SPS 문서 생성 실행 리포트",
            "",
            f"- Request ID: {request_id}",
            f"- Requested By: {requested_by}",
            f"- Generated: {generated_dt}",
            f"- DOCX: {docx_path.resolve()}",
            f"- Document blocks: {len(document.blocks)}",
            "",
            "## 자동 번호 생성 결과",
            "",
            *[f"- {heading}" for heading in headings],
            "",
            "## 검증 결과",
            "",
            "- [x] 사용자 로그인 세션 확인",
            "- [x] 사용자 입력 Document Model 생성",
            "- [x] Part/Chapter/Section 계층 번호 생성",
            "- [x] DOCX 산출물 저장",
            "- [x] DOCX Heading 텍스트 일치",
            "- [ ] Microsoft Word 육안 점검 필요",
            "",
        ]
        report_path.write_text("\n".join(report_lines), encoding="utf-8")
        return report_path
