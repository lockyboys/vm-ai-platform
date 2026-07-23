"""Generate the SPS official document-generation demonstration document."""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document

from engine.processor.word.document_model import (
    DocumentHeading,
    DocumentMetadata,
    DocumentModel,
)
from engine.processor.word.word_writer import WordWriter


def build_demo_document() -> DocumentModel:
    """Build an official SPS document-generation demonstration document."""

    model = DocumentModel(
        title="Story Programming Framework 문서 생성 표준 v1.0",
        metadata=DocumentMetadata(
            document_name="sps_document_generation_standard_v1_0",
            source_object_code="DOCUMENT",
            description=(
                "Story Programming Framework 공식 문서 구조를 생성하고 "
                "DOCX 산출물과 실행 검증 리포트를 함께 생성한다."
            ),
        ),
    )

    model.add_numbered_heading(
        "Story Programming Framework 문서 생성 원칙",
        level=1,
        heading_label="Part",
    )
    model.add_paragraph(
        "사람이 Story를 정의하고, 저장소(Repository)가 Story를 기억하며, "
        "Engine(엔진)이 Story를 이해하고, 생성기(Generator)가 Story를 구현한다.",
        italic=True,
    )

    model.add_numbered_heading(
        "문서 모델과 생성 흐름",
        level=2,
        heading_label="Chapter",
    )
    model.add_numbered_heading(
        "문서 구조",
        level=3,
    )
    model.add_paragraph(
        "문서는 Part, Chapter, Section, Paragraph의 계층으로 구성한다."
    )
    model.add_paragraph(
        "문서 구조와 산출 형식은 중립 Document Model에 기록하고, "
        "Word Generator는 해당 모델을 DOCX로 변환한다."
    )

    model.add_numbered_heading(
        "자동 번호 생성",
        level=3,
    )
    model.add_paragraph(
        "제목 레벨에 따라 Part 1., Chapter 1.1., Section 1.1.1. 형식의 "
        "계층 번호를 자동 생성한다."
    )
    model.add_paragraph(
        "상위 제목이 추가되면 하위 제목 번호는 초기화하며, 문서 구조는 "
        "Generator가 재사용할 수 있는 공식 Metadata로 유지한다."
    )

    model.add_numbered_heading(
        "DOCX 생성 및 검증",
        level=2,
        heading_label="Chapter",
    )
    model.add_paragraph(
        "생성된 DocumentModel은 Microsoft Word에서 열 수 있는 DOCX 파일로 저장한다."
    )
    model.add_paragraph(
        "실행 리포트는 자동 번호 결과, DOCX Heading 구조, 산출물 경로 및 "
        "Microsoft Word 육안 점검 상태를 기록한다."
    )

    return model


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="SPS 공식 문서 생성 DOCX 시연을 실행합니다.",
    )
    parser.add_argument(
        "--output",
        default="output/sps_document_generation_standard_v1_0.docx",
        help="생성할 DOCX 파일 경로",
    )
    return parser.parse_args()


def read_exported_headings(output_path: Path) -> list[str]:
    """저장된 DOCX에서 제목 텍스트를 다시 읽어 구조를 검증한다."""

    exported_document = Document(output_path)
    return [
        paragraph.text.strip()
        for paragraph in exported_document.paragraphs
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip()
    ]


def write_demo_report(
    *,
    document: DocumentModel,
    output_path: Path,
    exported_headings: list[str],
) -> Path:
    """공식 문서 생성 결과와 검증 상태를 Markdown 리포트로 저장한다."""

    report_path = output_path.with_name(
        f"{output_path.stem}_report.md"
    )
    numbered_headings = [
        block.text
        for block in document.blocks
        if isinstance(block, DocumentHeading)
    ]
    heading_validation_passed = exported_headings == numbered_headings
    official_structure_passed = numbered_headings == [
        "Part 1. Story Programming Framework 문서 생성 원칙",
        "Chapter 1.1. 문서 모델과 생성 흐름",
        "1.1.1. 문서 구조",
        "1.1.2. 자동 번호 생성",
        "Chapter 1.2. DOCX 생성 및 검증",
    ]
    test_label_absent = all(
        forbidden_text not in "\n".join(numbered_headings)
        for forbidden_text in ("Demo", "Part:", "Chapter:", "Section:")
    )

    report_lines = [
        "# SPS 문서 생성 실행 리포트",
        "",
        "- Document Type: Story Programming Framework 공식 기술 표준",
        "- Version: v1.0",
        (
            "- Generated: "
            f"{datetime.now(ZoneInfo('Asia/Seoul')).strftime('%Y-%m-%d %H:%M:%S KST')}"
        ),
        f"- DOCX: {output_path.resolve()}",
        f"- Document blocks: {len(document.blocks)}",
        "",
        "## 자동 번호 생성 결과",
        "",
        *[f"- {heading}" for heading in numbered_headings],
        "",
        "## DOCX 구조 검증",
        "",
        *[f"- {heading}" for heading in exported_headings],
        "",
        "## 검증 결과",
        "",
        "- [x] 공식 문서 모델 생성",
        "- [x] Part, Chapter, Section, Paragraph 구조 생성",
        "- [x] Part/Chapter/Section 계층 번호 생성",
        "- [x] DOCX 산출물 저장",
        (
            "- [x] DOCX Heading 텍스트 일치"
            if heading_validation_passed
            else "- [ ] DOCX Heading 텍스트 불일치"
        ),
        (
            "- [x] 공식 문서 제목 계층 검증"
            if official_structure_passed
            else "- [ ] 공식 문서 제목 계층 검증 실패"
        ),
        (
            "- [x] 테스트용 라벨 제거 검증"
            if test_label_absent
            else "- [ ] 테스트용 라벨 제거 검증 실패"
        ),
        "- [ ] Microsoft Word 육안 점검 필요",
        "",
    ]
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(report_lines), encoding="utf-8")
    return report_path.resolve()


def main() -> None:
    arguments = parse_arguments()
    output_path = Path(arguments.output)

    document = build_demo_document()
    saved_path = WordWriter().save(
        model=document,
        output_path=output_path,
    )
    exported_headings = read_exported_headings(saved_path)
    report_path = write_demo_report(
        document=document,
        output_path=saved_path,
        exported_headings=exported_headings,
    )

    print("SPS Document Generation SUCCESS")
    print(f"Output : {saved_path}")
    print(f"Report : {report_path}")


if __name__ == "__main__":
    main()
